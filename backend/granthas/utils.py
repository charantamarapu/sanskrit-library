from docx import Document
from io import BytesIO
import time
from copy import deepcopy


def get_paragraph_heading_level(paragraph):
    """
    Get heading level ONLY from Word's built-in Heading styles.
    Returns 1-9 if heading, None if body text.
    """
    style_name = paragraph.style.name
    
    if style_name.startswith('Heading'):
        try:
            level_str = style_name.replace('Heading', '').strip()
            if level_str.isdigit():
                level = int(level_str)
                if 1 <= level <= 9:
                    return level
            elif level_str == '':
                return 1
        except:
            pass
    
    return None


def filter_docx_by_commentaries(docx_path, selected_commentaries):
    """Filter Word document by selected commentaries."""
    
    start_time = time.time()
    print(f"\n{'='*70}")
    print("FILTERING STARTED")
    print(f"{'='*70}\n")
    
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"ERROR loading document: {e}")
        raise
    
    # Create new document with copied styles
    new_doc = Document()
    new_doc.styles._element = deepcopy(doc.styles._element)
    
    # Remove default empty paragraph
    if len(new_doc.paragraphs) > 0:
        p = new_doc.paragraphs[0]._element
        p.getparent().remove(p)
    
    all_commentaries = selected_commentaries.get('all_commentaries', [])
    selected = selected_commentaries.get('selected', [])
    
    include_all = 'all' in selected
    nothing_selected = len(selected) == 0
    
    skip_until_level = None
    included = 0
    excluded = 0
    commentary_actions = {}
    
    print(f"Commentaries: {all_commentaries}")
    print(f"Selected: {selected}")
    print(f"Mode: {'ALL' if include_all else 'NONE' if nothing_selected else 'SELECTIVE'}\n")
    
    for i, paragraph in enumerate(doc.paragraphs):
        if i > 0 and i % 1000 == 0:
            print(f"Processing: {i}/{len(doc.paragraphs)} paragraphs...")
        
        text = paragraph.text.strip()
        
        if not text:
            if skip_until_level is None:
                new_para_element = deepcopy(paragraph._element)
                new_doc._element.body.append(new_para_element)
                included += 1
            else:
                excluded += 1
            continue
        
        heading_level = get_paragraph_heading_level(paragraph)
        
        if heading_level is not None:
            if skip_until_level is not None:
                if heading_level <= skip_until_level:
                    skip_until_level = None
                else:
                    excluded += 1
                    continue
            
            if text in all_commentaries:
                if text not in commentary_actions:
                    if include_all:
                        print(f"✓ Level {heading_level}: '{text}' → INCLUDE (All)")
                        commentary_actions[text] = 'include'
                    elif nothing_selected:
                        print(f"✗ Level {heading_level}: '{text}' → REMOVE (None)")
                        commentary_actions[text] = 'remove'
                    elif text in selected:
                        print(f"✓ Level {heading_level}: '{text}' → INCLUDE (Selected)")
                        commentary_actions[text] = 'include'
                    else:
                        print(f"✗ Level {heading_level}: '{text}' → REMOVE (Not selected)")
                        commentary_actions[text] = 'remove'
                
                if include_all or (text in selected):
                    skip_until_level = None
                elif nothing_selected or (text not in selected):
                    skip_until_level = heading_level
                    excluded += 1
                    continue
        
        if skip_until_level is not None:
            excluded += 1
            continue
        
        new_para_element = deepcopy(paragraph._element)
        new_doc._element.body.append(new_para_element)
        included += 1
    
    elapsed = time.time() - start_time
    
    print(f"\n{'='*70}")
    print(f"COMPLETED:")
    print(f"  Total: {len(doc.paragraphs)}, Included: {included}, Excluded: {excluded}")
    print(f"  Time: {elapsed:.2f}s")
    print(f"{'='*70}\n")
    
    buffer = BytesIO()
    new_doc.save(buffer)
    buffer.seek(0)
    
    return buffer
